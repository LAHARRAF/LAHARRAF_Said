"""
Generate professional Word (.docx) CVs for Said LAHARRAF in French and English,
matching the exact single-page academic/professional layout shown in the reference image.

Requirements:
    pip install python-docx Pillow

Usage:
    python generate_cv.py
"""

import os
from PIL import Image, ImageDraw, ImageOps
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ──────────────────────────────────────────────────────────────
# HELPERS FOR XML FORMATTING & STYLING
# ──────────────────────────────────────────────────────────────

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set zero or customized padding (margins) on a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(int(val * 20)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def make_circle_image(img_path, output_path):
    """Crops the profile image to a perfect circle and saves it as a transparent PNG."""
    img = Image.open(img_path).convert("RGBA")
    size = min(img.size)
    # Crop to a square center
    img = ImageOps.fit(img, (size, size), Image.Resampling.LANCZOS)
    
    # Create circular mask
    mask = Image.new("L", (size, size), 0)
    draw = ImageDraw.Draw(mask)
    draw.ellipse((0, 0, size, size), fill=255)
    
    # Apply mask
    output = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    output.paste(img, (0, 0), mask=mask)
    output.save(output_path, "PNG")


def add_hyperlink(paragraph, text, url, color="a9326f", underline=True):
    """Adds a clickable hyperlink to a paragraph with Times New Roman font."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)

    # Set Size (9.5 pt)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '19')  # 19 half-points = 9.5 pt
    rPr.append(sz)
    
    # Set Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_section_heading(doc, text):
    """Adds a section heading preceded by a thin black horizontal line."""
    # Add a thin line paragraph
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(2)
    p_line.paragraph_format.keep_with_next = True
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 6/8 pt size
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Add centered heading text
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(4)
    p_title.paragraph_format.keep_with_next = True
    run = p_title.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_experience_entry(doc, company, role, loc_date, bullets):
    """Adds a standard experience entry: Company, followed by Role (left) & Loc/Date (right), then bullets."""
    # Company Name line (simple paragraph, bold, left-aligned)
    p_comp = doc.add_paragraph()
    p_comp.paragraph_format.space_before = Pt(3)
    p_comp.paragraph_format.space_after = Pt(0)
    p_comp.paragraph_format.keep_with_next = True
    run_company = p_comp.add_run(company)
    run_company.font.name = "Times New Roman"
    run_company.font.size = Pt(10)
    run_company.font.bold = True
    run_company.font.italic = True
    
    # Role & Date line (table with 2 columns, borderless)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.8)
    table.columns[1].width = Cm(4.8)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    set_cell_margins(cell_left, top=0, bottom=0, left=0, right=0)
    set_cell_margins(cell_right, top=0, bottom=0, left=0, right=0)
    
    # Left: Role Name (bold + italic)
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(0)
    p_left.paragraph_format.space_after = Pt(1.5)
    p_left.paragraph_format.line_spacing = 1.0
    p_left.paragraph_format.keep_with_next = True
    run_role = p_left.add_run(role)
    run_role.font.name = "Times New Roman"
    run_role.font.size = Pt(9.5)
    run_role.font.bold = True
    
    # Right: Location & Date (bold, right-aligned)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_before = Pt(0)
    p_right.paragraph_format.space_after = Pt(1.5)
    p_right.paragraph_format.line_spacing = 1.0
    run_loc = p_right.add_run(loc_date)
    run_loc.font.name = "Times New Roman"
    run_loc.font.size = Pt(9.5)
    run_loc.font.bold = True

    # Bullets
    for bullet in bullets:
        add_bullet_point(doc, bullet)


def add_bullet_point(doc, text):
    """Adds a tight bullet point starting with a dash, with bold colon prefixes formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.05
    
    # Bullet symbol
    run_bullet = p.add_run("-   ")
    run_bullet.font.name = "Times New Roman"
    run_bullet.font.size = Pt(9.5)
    
    # Format prefix if contains colon
    if " : " in text and text.index(" : ") < 30:
        prefix, rest = text.split(" : ", 1)
        run_pref = p.add_run(prefix + " : ")
        run_pref.font.name = "Times New Roman"
        run_pref.font.size = Pt(9.5)
        run_pref.bold = True
        
        run_rest = p.add_run(rest)
        run_rest.font.name = "Times New Roman"
        run_rest.font.size = Pt(9.5)
    elif ": " in text and text.index(": ") < 30:
        prefix, rest = text.split(":", 1)
        run_pref = p.add_run(prefix + " : ")
        run_pref.font.name = "Times New Roman"
        run_pref.font.size = Pt(9.5)
        run_pref.bold = True
        
        run_rest = p.add_run(rest)
        run_rest.font.name = "Times New Roman"
        run_rest.font.size = Pt(9.5)
    else:
        run_text = p.add_run(text)
        run_text.font.name = "Times New Roman"
        run_text.font.size = Pt(9.5)


def add_education_entry(doc, institution, loc_date, bullets):
    """Adds an education entry: Institution (left) and Loc/Date (right) on the same line, then bullets."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.8)
    table.columns[1].width = Cm(4.8)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    set_cell_margins(cell_left, top=0, bottom=0, left=0, right=0)
    set_cell_margins(cell_right, top=0, bottom=0, left=0, right=0)
    
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(3)
    p_left.paragraph_format.space_after = Pt(1.5)
    p_left.paragraph_format.keep_with_next = True
    run_inst = p_left.add_run(institution)
    run_inst.font.name = "Times New Roman"
    run_inst.font.size = Pt(10)
    run_inst.font.bold = True
    
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_before = Pt(3)
    p_right.paragraph_format.space_after = Pt(1.5)
    run_loc = p_right.add_run(loc_date)
    run_loc.font.name = "Times New Roman"
    run_loc.font.size = Pt(9.5)
    run_loc.font.bold = True

    for bullet in bullets:
        add_bullet_point(doc, bullet)


def add_single_line_entry(doc, left_text, right_text):
    """Adds a single row entry spanning left and right margins (projects/extracurriculars)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.8)
    table.columns[1].width = Cm(4.8)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    set_cell_margins(cell_left, top=0, bottom=0, left=0, right=0)
    set_cell_margins(cell_right, top=0, bottom=0, left=0, right=0)
    
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(1)
    p_left.paragraph_format.space_after = Pt(1.5)
    
    # Bullet text formatting
    if ": " in left_text and left_text.index(": ") < 60:
        bold_part, rest = left_text.split(": ", 1)
        run_b = p_left.add_run(bold_part + " : ")
        run_b.bold = True
        run_b.font.name = "Times New Roman"
        run_b.font.size = Pt(9.5)
        
        run_r = p_left.add_run(rest)
        run_r.font.name = "Times New Roman"
        run_r.font.size = Pt(9.5)
    else:
        run_l = p_left.add_run(left_text)
        run_l.font.name = "Times New Roman"
        run_l.font.size = Pt(9.5)
        
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_before = Pt(1)
    p_right.paragraph_format.space_after = Pt(1.5)
    run_r = p_right.add_run(right_text)
    run_r.font.name = "Times New Roman"
    run_r.font.size = Pt(9.5)
    run_r.font.bold = True


def add_skills_entry_parsed(doc, text):
    """Adds a competence entry and bolds multiple 'Key :' sections dynamically."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.05
    
    # Bullet symbol
    run_bullet = p.add_run("-   ")
    run_bullet.font.name = "Times New Roman"
    run_bullet.font.size = Pt(9.5)
    
    # Split by " | "
    parts = text.split(" | ")
    for idx, part in enumerate(parts):
        if idx > 0:
            run_sep = p.add_run(" | ")
            run_sep.font.name = "Times New Roman"
            run_sep.font.size = Pt(9.5)
            
        # Split key and value
        if " : " in part:
            key, val = part.split(" : ", 1)
            run_k = p.add_run(key + " : ")
            run_k.font.name = "Times New Roman"
            run_k.font.size = Pt(9.5)
            run_k.bold = True
            
            run_v = p.add_run(val)
            run_v.font.name = "Times New Roman"
            run_v.font.size = Pt(9.5)
        elif ":" in part:
            key, val = part.split(":", 1)
            run_k = p.add_run(key + " : ")
            run_k.font.name = "Times New Roman"
            run_k.font.size = Pt(9.5)
            run_k.bold = True
            
            run_v = p.add_run(val)
            run_v.font.name = "Times New Roman"
            run_v.font.size = Pt(9.5)
        else:
            run_p = p.add_run(part)
            run_p.font.name = "Times New Roman"
            run_p.font.size = Pt(9.5)


def add_certifications_section(doc, left_title, left_certs, right_title, right_certs, section_title):
    """Adds a 2-column layout for certifications."""
    add_section_heading(doc, section_title)
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(9.3)
    table.columns[1].width = Cm(9.3)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    set_cell_margins(cell_left, top=0, bottom=0, left=0, right=0)
    set_cell_margins(cell_right, top=0, bottom=0, left=0, right=0)
    
    # Left column: DataCamp
    p_left_title = cell_left.paragraphs[0]
    p_left_title.paragraph_format.space_before = Pt(3)
    p_left_title.paragraph_format.space_after = Pt(2)
    run_lt = p_left_title.add_run(left_title + " :")
    run_lt.font.name = "Times New Roman"
    run_lt.font.size = Pt(10)
    run_lt.font.bold = True
    
    for cert in left_certs:
        p = cell_left.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        run = p.add_run("-   " + cert)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.5)
        
    # Right column: Oracle
    p_right_title = cell_right.paragraphs[0]
    p_right_title.paragraph_format.space_before = Pt(3)
    p_right_title.paragraph_format.space_after = Pt(2)
    run_rt = p_right_title.add_run(right_title + " :")
    run_rt.font.name = "Times New Roman"
    run_rt.font.size = Pt(10)
    run_rt.font.bold = True
    
    for cert in right_certs:
        p = cell_right.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        run = p.add_run("-   " + cert)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.5)


# ──────────────────────────────────────────────────────────────
# DATA DECLARATIONS
# ──────────────────────────────────────────────────────────────

CV_DATA_FR = {
    "sections": {
        "experience": "Expérience Professionnelle",
        "projects": "Projets Académiques (Data & IA)",
        "formation": "Formation",
        "skills": "Compétences & Centres d'intérêt",
        "certifications": "Certifications",
        "extracurricular": "Activités Para-universitaires"
    },
    "experience": [
        {
            "company": "Université Al Akhawayn (AUI)",
            "role": "Ingénieur Data : Concepteur de base de données – Projet de fin d'études",
            "loc_date": "Ifrane, Novembre - Juin 2025",
            "bullets": [
                "Conception et mise en place d'une plateforme de suivi des étudiants-athlètes depuis zéro.",
                "Analyse et modélisation des processus (BPMN) pour représenter l'état actuel.",
                "Conception de modèles UML (diagrammes de cas d'utilisation, de classes).",
                "Conception, développement et gestion de bases SQL pour le stockage et l'analyse de performance.",
                "Mise en place de pipelines ETL pour l'intégration et le suivi des données.",
                "Collaboration avec les parties prenantes pour la collecte des besoins et la documentation.",
                "Mise en place d'une architecture scalable sur GCP pour héberger la plateforme et déployer des modèles ML.",
                "Technos : MySQL, SQL, UML, GCP, BPMN, phpMyAdmin, Lucidchart."
            ]
        },
        {
            "company": "Freelance",
            "role": "Développeur Full-Stack Web & Mobile",
            "loc_date": "Remote, 2024 – Présent",
            "bullets": [
                "Conception de l'architecture d'applications web, mobiles et de bureau, du concept au déploiement en production.",
                "Clinique Dentaire : Développement, gestion et déploiement d'un système de gestion de cabinet médical complet pour optimiser les dossiers patients et la planification.",
                "Magasin de Pneus : Création d'une plateforme e-commerce et de gestion des stocks (du catalogue produit à la validation des commandes).",
                "Conception de bases de données relationnelles SQL et administration système avec déploiement sur serveurs d'hébergement via cPanel.",
                "Technos : PHP, Laravel, Python, MySQL, HTML, CSS, JavaScript, cPanel."
            ]
        },
        {
            "company": "Arteka :",
            "role": "Data Analyste – Stage d'été",
            "loc_date": "France (télétravail), Août – Octobre 2024",
            "bullets": [
                "Extraction et préparation de jeux de données sportifs (CSV/JSON, APIs en ligne).",
                "Prétraitement et structuration des données pour l'entraînement de modèles ML.",
                "Création de dashboards interactifs pour l'analyse de matchs de football et tennis.",
                "Benchmarking de plateformes d'analytique sportive.",
                "Technos : Python (pandas, matplotlib, mplsoccer, NumPy), LabelMe."
            ]
        },
        {
            "company": "Physicare :",
            "role": "Développeur Mobile – Projet de fin d'études",
            "loc_date": "France (télétravail), Avril – Juin 2023",
            "bullets": [
                "Conception d'une application mobile pour le suivi du diabète et la prescription sportive.",
                "Réalisation de maquettes UI et développement des fonctionnalités principales.",
                "Technos : Flutter, Firebase, Adobe XD."
            ]
        },
        {
            "company": "Raznet :",
            "role": "Développeur Full-Stack – Stage d'été",
            "loc_date": "Casablanca, Août - Octobre 2022",
            "bullets": [
                "Conception d'architecture fonctionnelle et schéma de base de données pour une application web de gestion de cabinet dentaire.",
                "Développement front-end et back-end avec sécurisation des données et UX fluide.",
                "Technos : PHP, MySQL, XAMPP, HTML, CSS, JS."
            ]
        },
        {
            "company": "Radeec :",
            "role": "Développeur Full-Stack – Projet de fin d'études",
            "loc_date": "Settat, Mai - Juin 2022",
            "bullets": [
                "Conception de l'architecture et du modèle de données d'une application web de gestion des bénéficiaires.",
                "Développement d'interfaces utilisateurs et fonctionnalités de gestion (création, suivi, reporting).",
                "Technos : PHP, XAMPP, phpMyAdmin, MySQL."
            ]
        }
    ],
    "projects": [
        ("- Big Data sur la Premier League (Hadoop, Hive, Spark, Python) : Collecte, analyse et visualisation de données.", "2024 – 2025"),
        ("- Étude statistique de l'activité physique (SPSS, PCA, XLSTAT) : Corrélation entre paramètres corporels et performance sportive.", "2023 – 2024"),
        ("- Analyse de performance au sprint (Python, scikit-learn, NumPy) : Étude quantitative des indicateurs de vitesse et force musculaire.", "2023 – 2024"),
        ("- Détection joueurs & ballon (YOLOv5, Python, LabelMe) : Création dataset et entraînement d'un modèle de détection pour l'analyse sportive.", "2024 – 2025"),
        ("- Système personnalisé de suivi sportif (PHP, MySQL, Bootstrap) : Recommandations d'entraînement basées sur tests physiques.", "2023 – 2024")
    ],
    "formation": [
        {
            "institution": "Université Hassan 1er - ISSS – Settat",
            "loc_date": "Settat, 2025 – Present",
            "bullets": [
                "Doctorat en intelligence artificielle et santé numérique.",
                "Assistant d'enseignement : Informatique Biomédicale."
            ]
        },
        {
            "institution": "Université Hassan 1er - I2S – Settat",
            "loc_date": "Settat, 2022 – 2025",
            "bullets": [
                "Master en Transformation Digitale et Technologies du Sport.",
                "Licence Professionnelle en Sport et Technologie."
            ]
        },
        {
            "institution": "Centre BTS",
            "loc_date": "Settat, 2020 – 2022",
            "bullets": [
                "Brevet de Technicien Supérieur (BTS) en Développement des Systèmes d'Information."
            ]
        },
        {
            "institution": "Lycée Technique Qualifiant",
            "loc_date": "Settat, 2019 – 2020",
            "bullets": [
                "Baccalauréat Sciences et Technologies Électriques (STE)."
            ]
        }
    ],
    "skills": [
        "Data Engineering : SQL, MySQL, Hadoop, Hive, Spark, Modélisation, ETL (Airflow), data warehouse (Snowflake) | Méthodologie : Agile, Scrum, Kanban",
        "Développement logiciel : PHP, LARAVEL, HTML, CSS, JavaScript, Flutter, C, JAVA",
        "Data Analytics & ML : Python (pandas, NumPy, scikit-learn, matplotlib, YOLO), SPSS, XLSTAT, Power BI, Tableau, Excel",
        "Modélisation & Outils : UML, BPMN, Merise, Git, VSCode, LabelMe, Firebase | Cloud : GCP, AWS",
        "Intérêts : Course à pied, Natation, Randonnée | Languages : Arabe (natif), Français (courant), Anglais (courant)"
    ],
    "certifications": {
        "datacamp": [
            "Data Engineer Associate",
            "Understanding Artificial Intelligence",
            "AWS Cloud Practitioner (CLF-C02)"
        ],
        "oracle": [
            "Oracle Cloud Infrastructure AI Foundations Associate (2025)",
            "Oracle Cloud Infrastructure AI Foundations Associate (2023)",
            "Oracle MySQL Implementation Associate"
        ]
    },
    "extracurricular": [
        ("- Participation à la 14ème édition du Bootcamp UM6P Entrepreneur Génération", "UM6P Benguerir, Juin. 2024"),
        ("- Présentation au Congrès International sur les Sciences de l'Activité Physique et Technologies Émergentes (ICPASET'24)", "UH1 Settat, Mai. 2024"),
        ("- Participation au Hackathon International Smart Cities", "ENS Martil, Mai. 2024"),
        ("- Participation au Hackathon Mobilité Grande Vitesse ONCF-UM6P", "UM6P Benguerir, Mars. 2024"),
        ("- Participation au NEXT-GEN Hackathon", "GDSC EMSI Rabat, Août. 2023"),
        ("- Participation au RACATHON 3.0 at EL Jadida", "ENSA El Jadida, Mai. 2023")
    ]
}

CV_DATA_EN = {
    "sections": {
        "experience": "Professional Experience",
        "projects": "Academic Projects (Data & AI)",
        "formation": "Education",
        "skills": "Skills & Interests",
        "certifications": "Certifications",
        "extracurricular": "Extracurricular Activities"
    },
    "experience": [
        {
            "company": "Al Akhawayn University (AUI)",
            "role": "Data Engineer : Database Designer – Final Year Project",
            "loc_date": "Ifrane, November - June 2025",
            "bullets": [
                "Design and implementation of a student-athlete tracking platform from scratch.",
                "Process analysis and modeling (BPMN) to represent the current state.",
                "UML modeling (use case and class diagrams).",
                "Design, development, and management of SQL databases for performance storage and analysis.",
                "ETL pipeline implementation for data integration and tracking.",
                "Collaboration with stakeholders for requirements gathering and documentation.",
                "Implementation of a scalable architecture on GCP to host the platform and deploy ML models.",
                "Technos : MySQL, SQL, UML, GCP, BPMN, phpMyAdmin, Lucidchart."
            ]
        },
        {
            "company": "Freelance",
            "role": "Full-Stack Web & Mobile Developer",
            "loc_date": "Remote, 2024 – Present",
            "bullets": [
                "Design application architectures across web, mobile, and desktop environments, driving projects from concept to production.",
                "Dental Clinic Platform : Developed, managed, and deployed a comprehensive management system to optimize medical records and patient scheduling.",
                "Tire Shop & E-Commerce : Created a full-process inventory and sales system managing product visualization, online ordering, and validation.",
                "Relational database schema design (SQL) and system administration including deployment on production hosting environments via cPanel.",
                "Technos : PHP, Laravel, Python, MySQL, HTML, CSS, JavaScript, cPanel."
            ]
        },
        {
            "company": "Arteka :",
            "role": "Data Analyst – Summer Internship",
            "loc_date": "France (remote), August – October 2024",
            "bullets": [
                "Extraction and preparation of sports datasets (CSV/JSON, online APIs).",
                "Data preprocessing and structuring for ML model training.",
                "Creation of interactive dashboards for football and tennis match analysis.",
                "Benchmarking of sports analytics platforms.",
                "Technos : Python (pandas, matplotlib, mplsoccer, NumPy), LabelMe."
            ]
        },
        {
            "company": "Physicare :",
            "role": "Mobile Developer – Graduation Project",
            "loc_date": "France (remote), April – June 2023",
            "bullets": [
                "Design of mobile application for diabetes tracking and sports prescription.",
                "UI mockup creation and development of main features.",
                "Technos : Flutter, Firebase, Adobe XD."
            ]
        },
        {
            "company": "Raznet :",
            "role": "Full-Stack Developer – Summer Internship",
            "loc_date": "Casablanca, August - October 2022",
            "bullets": [
                "Functional architecture and database schema design for dental clinic management web application.",
                "Front-end and back-end development with data security and smooth UX.",
                "Technos : PHP, MySQL, XAMPP, HTML, CSS, JS."
            ]
        },
        {
            "company": "Radeec :",
            "role": "Full-Stack Developer – Graduation Project",
            "loc_date": "Settat, May - June 2022",
            "bullets": [
                "Architecture and data model design for beneficiary management web application.",
                "User interface development and management features (creation, tracking, reporting).",
                "Technos : PHP, XAMPP, phpMyAdmin, MySQL."
            ]
        }
    ],
    "projects": [
        ("- Big Data on the Premier League (Hadoop, Hive, Spark, Python) : Data collection, analysis and visualization.", "2024 – 2025"),
        ("- Physical Activity Statistical Study (SPSS, PCA, XLSTAT) : Correlation between body parameters and sports performance.", "2023 – 2024"),
        ("- Sprint Performance Analysis (Python, scikit-learn, NumPy) : Quantitative study of speed and muscular strength indicators.", "2023 – 2024"),
        ("- Player & Ball Detection (YOLOv5, Python, LabelMe) : Custom dataset creation and YOLOv5 model training for sports analytics.", "2024 – 2025"),
        ("- Personalized Sports Tracking System (PHP, MySQL, Bootstrap) : Training recommendations based on physical fitness tests.", "2023 – 2024")
    ],
    "formation": [
        {
            "institution": "Hassan First University - ISSS – Settat",
            "loc_date": "Settat, 2025 – Present",
            "bullets": [
                "PhD Candidate in Artificial Intelligence & Digital Health.",
                "Teaching Assistant: Biomedical Informatics."
            ]
        },
        {
            "institution": "Hassan First University - I2S – Settat",
            "loc_date": "Settat, 2022 – 2025",
            "bullets": [
                "Master's in Digital Transformation and Sports Technologies.",
                "Professional Bachelor's in Sport and Technology."
            ]
        },
        {
            "institution": "BTS Center",
            "loc_date": "Settat, 2020 – 2022",
            "bullets": [
                "Higher Technician Certificate (BTS) in Information Systems Development."
            ]
        },
        {
            "institution": "Technical Qualifying High School",
            "loc_date": "Settat, 2019 – 2020",
            "bullets": [
                "Baccalaureate in Electrical Sciences and Technologies (STE)."
            ]
        }
    ],
    "skills": [
        "Data Engineering : SQL, MySQL, Hadoop, Hive, Spark, Data Modeling, ETL (Airflow), data warehouse (Snowflake) | Methodology : Agile, Scrum, Kanban",
        "Software Development : PHP, LARAVEL, HTML, CSS, JavaScript, Flutter, C, JAVA",
        "Data Analytics & ML : Python (pandas, NumPy, scikit-learn, matplotlib, YOLO), SPSS, XLSTAT, Power BI, Tableau, Excel",
        "Modeling & Tools : UML, BPMN, Merise, Git, VSCode, LabelMe, Firebase | Cloud : GCP, AWS",
        "Interests : Running, Swimming, Hiking | Languages : Arabic (native), French (fluent), English (fluent)"
    ],
    "certifications": {
        "datacamp": [
            "Data Engineer Associate",
            "Understanding Artificial Intelligence",
            "AWS Cloud Practitioner (CLF-C02)"
        ],
        "oracle": [
            "Oracle Cloud Infrastructure AI Foundations Associate (2025)",
            "Oracle Cloud Infrastructure AI Foundations Associate (2023)",
            "Oracle MySQL Implementation Associate"
        ]
    },
    "extracurricular": [
        ("- Participation in the 14th edition of the UM6P Entrepreneur Generation Bootcamp", "UM6P Benguerir, June 2024"),
        ("- Presentation at the International Congress on Physical Activity Sciences and Emerging Technologies (ICPASET'24)", "UH1 Settat, May 2024"),
        ("- Participation in the Smart Cities International Hackathon", "ENS Martil, May 2024"),
        ("- Participation in the ONCF-UM6P High-Speed Mobility Hackathon", "UM6P Benguerir, March 2024"),
        ("- Participation in the NEXT-GEN Hackathon", "GDSC EMSI Rabat, August 2023"),
        ("- Participation in the RACATHON 3.0 at El Jadida", "ENSA El Jadida, May 2023")
    ]
}


# ──────────────────────────────────────────────────────────────
# GENERATE SINGLE DOCUMENT
# ──────────────────────────────────────────────────────────────

def generate_document(data, filename, is_french=True):
    doc = Document()

    # ── Page margins ──
    # Tight margins (1.0 cm top/bottom, 1.2 cm left/right) to guarantee single-page fit
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(0, 0, 0)

    # ══════════════════════════════════════════════════════════
    # HEADER (Name, details on left, circular photo on right)
    # ══════════════════════════════════════════════════════════
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.0)
    table.columns[1].width = Cm(3.6)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    set_cell_margins(cell_left, top=0, bottom=0, left=0, right=0)
    set_cell_margins(cell_right, top=0, bottom=0, left=0, right=0)
    
    # Left Cell: Name & Contact Details
    p_name = cell_left.paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("Said LAHARRAF")
    run_name.font.name = "Times New Roman"
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0, 0, 0)
    
    p_contact = cell_left.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(2)
    contact_text = "laharrafsaid@gmail.com | (+212) 7 77 60 35 26 | Casablanca-Settat, Morocco"
    run_contact = p_contact.add_run(contact_text)
    run_contact.font.name = "Times New Roman"
    run_contact.font.size = Pt(9.5)
    
    p_links = cell_left.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_links.paragraph_format.space_before = Pt(0)
    p_links.paragraph_format.space_after = Pt(0)
    
    # Add links (red-purple color matching the reference theme)
    add_hyperlink(p_links, "LinkedIn", "https://linkedin.com/in/said-laharraf", color="a9326f", underline=True)
    run_sep = p_links.add_run(", ")
    run_sep.font.name = "Times New Roman"
    run_sep.font.size = Pt(9.5)
    add_hyperlink(p_links, "Portfolio", "https://laharrafsaid.github.io", color="a9326f", underline=True)
    
    # Right Cell: Photo
    p_photo = cell_right.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)
    
    me_img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "images", "me.png")
    temp_circle_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp_me_circle.png")
    
    if os.path.exists(me_img_path):
        try:
            make_circle_image(me_img_path, temp_circle_path)
            run_photo = p_photo.add_run()
            run_photo.add_picture(temp_circle_path, width=Cm(2.5))
        except Exception as e:
            print(f"Error processing circular image: {e}")
            run_photo = p_photo.add_run()
            run_photo.add_picture(me_img_path, width=Cm(2.5))
    else:
        print(f"Warning: profile photo not found at {me_img_path}")

    # ══════════════════════════════════════════════════════════
    # EXPERIENCES
    # ══════════════════════════════════════════════════════════
    add_section_heading(doc, data["sections"]["experience"])
    for exp in data["experience"]:
        add_experience_entry(doc, exp["company"], exp["role"], exp["loc_date"], exp["bullets"])

    # ══════════════════════════════════════════════════════════
    # PROJECTS
    # ══════════════════════════════════════════════════════════
    add_section_heading(doc, data["sections"]["projects"])
    for proj in data["projects"]:
        add_single_line_entry(doc, proj[0], proj[1])

    # ══════════════════════════════════════════════════════════
    # EDUCATION
    # ══════════════════════════════════════════════════════════
    add_section_heading(doc, data["sections"]["formation"])
    for edu in data["formation"]:
        add_education_entry(doc, edu["institution"], edu["loc_date"], edu["bullets"])

    # ══════════════════════════════════════════════════════════
    # SKILLS & INTERESTS
    # ══════════════════════════════════════════════════════════
    add_section_heading(doc, data["sections"]["skills"])
    for skill in data["skills"]:
        add_skills_entry_parsed(doc, skill)

    # ══════════════════════════════════════════════════════════
    # CERTIFICATIONS (2 COLUMNS)
    # ══════════════════════════════════════════════════════════
    add_certifications_section(
        doc, 
        "DataCamp", 
        data["certifications"]["datacamp"], 
        "Oracle", 
        data["certifications"]["oracle"], 
        data["sections"]["certifications"]
    )

    # ══════════════════════════════════════════════════════════
    # EXTRACURRICULAR ACTIVITIES
    # ══════════════════════════════════════════════════════════
    add_section_heading(doc, data["sections"]["extracurricular"])
    for act in data["extracurricular"]:
        add_single_line_entry(doc, act[0], act[1])

    # Save output docx
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
    doc.save(output_path)
    print(f"[OK] Generated: {output_path}")
    
    # Cleanup temp photo
    if os.path.exists(temp_circle_path):
        try:
            os.remove(temp_circle_path)
        except:
            pass
            
    return output_path


# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    generate_document(CV_DATA_FR, "Said_LAHARRAF_CV_FR.docx", is_french=True)
    generate_document(CV_DATA_EN, "Said_LAHARRAF_CV_EN.docx", is_french=False)
